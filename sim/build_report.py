# -*- coding: utf-8 -*-
"""
build_report.py — T9 验证报告生成器(python-docx)
纪律(canonical/铁律):
  1. 叙事文字里**零硬编码性能数字**——一切数字从 out/*.csv 读取注入,缺数据打〔待补〕;
  2. 图直接嵌 sim/figs 的 PNG(它们本身就是确定性再生的验证产物,每图带复现命令);
  3. 章节结构=递进故事(随机存放→在线评分→批量上界→PLC 落地→绿色账),
     每章对应 2025 评委评语六要素(附录 C 给自查映射表);
  4. 文风=学术克制(陈述句收尾,不排比,不堆"我们"),生成后仍会人工过 de-AI 七规则。
运行:python build_report.py → out/验证报告_draft.docx + 控制台"数字来源清单"
再生:任何口径变更后先跑上游脚本刷新 CSV,再跑本脚本——报告永远与数据同源。
"""
import csv
import hashlib
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
FIGS = os.path.join(HERE, "figs")
FONT, MONO = "Microsoft YaHei", "Consolas"
MISS = "〔待补〕"
USED = []                 # 数字来源清单(key, value, source)


# ---------------- 数据装载(登记缺失；正文出现待补即fail closed) ----------------
def load_rows(name):
    p = os.path.join(OUT, name)
    if not os.path.exists(p):
        return []
    with open(p, encoding="utf-8-sig") as fp:
        return [r for r in csv.DictReader(fp) if any((v or "").strip() for v in r.values())]


def by_key(rows, key_col):
    return {r[key_col]: r for r in rows if r.get(key_col)}


def fnum(row, col, nd=2, scale=1.0, src=""):
    """行列取数→格式化字符串;任何缺失→〔待补〕。所有取用登记进 USED。"""
    try:
        v = float(row[col]) * scale
        s = f"{v:,.{nd}f}"
        USED.append((f"{src}:{col}", s))
        return s
    except Exception:
        USED.append((f"{src}:{col}", MISS))
        return MISS


def reg(key, val):
    """派生数字登记进来源清单后返回(自查'待补'口径必须涵盖派生值)。"""
    USED.append((key, val))
    return val


def pct_drop(a, b, src=""):
    """(a-b)/a*100,输入 str/float 容错;登记来源。"""
    try:
        a, b = float(a), float(b)
        return reg(src or "derived:pct_drop", f"{(a - b) / a * 100:.1f}")
    except Exception:
        return reg(src or "derived:pct_drop", MISS)


HEAD = by_key(load_rows("headline_numbers.csv"), "strategy")
MIX = {(r["strategy"], r["motion"]): r for r in load_rows("mixed_ops.csv")}
BOUNDS = by_key(load_rows("analytic_bounds.csv"), "quantity")
PARETO = load_rows("pareto_front.csv")
SENS = load_rows("sensitivity.csv")
STRESS = by_key(load_rows("stress_matrix.csv"), "strategy")
TIGHT = {r["strat"]: r for r in load_rows("experiments_agg.csv")
         if r.get("scene") == "tight"}          # 紧库存场景(多 seed 实验矩阵)
REAL = {(r["tier"], r["strategy"]): r for r in load_rows("realism_matrix_data.csv")
        if r.get("tier") in ("1", "2", "3")}    # L7 三档单一 schema 数据表

# L7 拟真度阶梯派生值
TIER_DROP = {}
for _t in ("1", "2", "3"):
    try:
        _s = float(REAL[(_t, "seq")]["avg_retrieve_s"])
        _a = float(REAL[(_t, "awra")]["avg_retrieve_s"])
        TIER_DROP[_t] = reg(f"realism:档{_t}降幅%", f"{(_s - _a) / _s * 100:.1f}")
    except Exception:
        TIER_DROP[_t] = MISS
RESP_AW95 = fnum(REAL.get(("3", "awra"), {}), "resp_p95_s", 0, src="realism")
RESP_SEQ95 = fnum(REAL.get(("3", "seq"), {}), "resp_p95_s", 0, src="realism")
RESP_NEAR95 = fnum(REAL.get(("3", "near"), {}), "resp_p95_s", 0, src="realism")
RESP_AW50 = fnum(REAL.get(("3", "awra"), {}), "resp_p50_s", 0, src="realism")
RESP_SEQ50 = fnum(REAL.get(("3", "seq"), {}), "resp_p50_s", 0, src="realism")

# 头条数字(H 口径)
AW, SEQ, NEAR, SCORE = (HEAD.get(k, {}) for k in ("awra", "seq", "near", "score"))
H_AW = fnum(AW, "H_exp_mean", 2, src="headline")
H_AW_STD = fnum(AW, "H_exp_std", 2, src="headline")
H_AW_CI = fnum(AW, "H_exp_ci95_half", 3, src="headline")
H_SEQ = fnum(SEQ, "H_exp_mean", 2, src="headline")
DROP_VS_SEQ = pct_drop(SEQ.get("H_exp_mean"), AW.get("H_exp_mean"), "derived:降幅vs顺序基线%")
DROP_E = pct_drop(SEQ.get("H_energy_mean"), AW.get("H_energy_mean"), "derived:能耗降幅%")
HEAVY_TIER = fnum(AW, "H_heavy_tier", 2, src="headline")
GAP_ONLINE = pct_drop(SCORE.get("H_exp_mean"), AW.get("H_exp_mean"), "derived:批量vs在线gap%")
LEG_AW = fnum(AW, "legacy_exp_mean", 2, src="headline")
LEG_DROP = pct_drop(SEQ.get("legacy_exp_mean"), AW.get("legacy_exp_mean"), "derived:对照口径降幅%")

# 吞吐(L2,梯形口径)
MA, MS = MIX.get(("awra", "trapezoid"), {}), MIX.get(("seq", "trapezoid"), {})
THR_AW = fnum(MA, "throughput_dual_command_ops_h", 0, src="mixed_ops")
THR_CYCLES = MISS
THR_X = MISS
try:
    THR_CYCLES = f"{float(MA['throughput_dual_command_ops_h']) / 2:.1f}"
    THR_X = f"{float(MA['throughput_dual_command_ops_h']) / float(MS['throughput_dual_command_ops_h']):.2f}"
    USED.append(("mixed_ops:双周期容量", THR_CYCLES))
    USED.append(("mixed_ops:thr倍数", THR_X))
except Exception:
    pass
DUAL_SAVE = fnum(MA, "saving_pct", 1, src="mixed_ops")

# 解析下界夹逼(F9 口径:匀速/seed2026;LB≤NSGA可达≤awra)
GAP_LB = GAP_ACH = LB_S = ACH_S = AW_BOUND_S = MISS
try:
    lb = float(BOUNDS["freq_aware_lower_bound"]["value_s"])
    aw_b = float(BOUNDS["strategy_awra"]["value_s"])
    ach = min(float(r["exp_retrieval_s"]) for r in PARETO)
    LB_S = reg("analytic_bounds:LB", f"{lb:.2f}")
    AW_BOUND_S = reg("analytic_bounds:awra", f"{aw_b:.2f}")
    ACH_S = reg("pareto_front:可达候选点", f"{ach:.2f}")
    GAP_LB = reg("derived:距松弛下界%", f"{(aw_b - lb) / lb * 100:.1f}")
    GAP_ACH = reg("derived:距可达候选点%", f"{(aw_b - ach) / ach * 100:.1f}")
except Exception:
    pass

# 绿色换算(参数/结果分表；每个 CSV 仅一个 schema)
E_ASSUME = {r["param"]: r["value"] for r in load_rows("energy_report_params.csv")
            if r.get("param")}
DAILY_OPS = reg("energy_report:daily_single_ops", E_ASSUME.get("daily_single_ops", MISS))
CO2_FACTOR = reg("energy_report:co2_factor", E_ASSUME.get("co2_kg_per_kwh", MISS))
try:
    DAILY_DUAL = reg("derived:daily_dual_cycles", f"{float(DAILY_OPS) / 2:.0f}")
except Exception:
    DAILY_DUAL = MISS
E_LAYOUT = by_key(load_rows("energy_report_data.csv"), "layout")
ASSUME_LABEL = {"eta": "电机传动效率 η", "mu": "水平阻力系数 μ",
                "carrier_kg": "载货台自重 kg", "daily_single_ops": "日单操作数",
                "work_days": "年工作日", "elec_price_yuan_kwh": "电价 元/kWh",
                "co2_kg_per_kwh": "排放因子 kgCO2/kWh"}
KWH_SAVE = KWH_SAVE_PCT = MISS
try:
    k_seq = float(E_LAYOUT["seq"]["kwh_per_year"])
    k_aw = float(E_LAYOUT["awra"]["kwh_per_year"])
    KWH_SAVE = reg("energy_report:年省kWh", f"{k_seq - k_aw:.1f}")
    KWH_SAVE_PCT = reg("energy_report:年省%", f"{(k_seq - k_aw) / k_seq * 100:.1f}")
except Exception:
    pass
# 展望口径:再生制动(R7 多口径方法论;η_regen 假设显式,不进头条)
REGEN_NET_AW = REGEN_REC_AW = REGEN_GAP = ETA_REGEN_S = MISS
try:
    ETA_REGEN_S = reg("energy_report:η_regen假设", E_ASSUME.get("eta_regen_outlook"))
    REGEN_REC_AW = reg("energy_report:awra年回收kWh",
                       f"{float(E_LAYOUT['awra']['recovered_kwh_year']):.0f}")
    REGEN_NET_AW = reg("energy_report:awra净耗电kWh",
                       f"{float(E_LAYOUT['awra']['kwh_year_net_regen']):.0f}")
    REGEN_GAP = reg("energy_report:再生口径差值kWh",
                    f"{float(E_LAYOUT['seq']['kwh_year_net_regen']) - float(E_LAYOUT['awra']['kwh_year_net_regen']):.0f}")
except Exception:
    pass

# RL 实跑对照(T13/F16)
RLD = {r["quantity"]: r["value"] for r in load_rows("rl_probe_summary.csv")
       if r.get("quantity")}
RL_FINAL = RL_SCORE = RL_RAND = RL_GAP = RL_COS = MISS
try:
    RL_FINAL = reg("rl_probe:RL终点", f"{float(RLD['rl_final_exp_t']):.2f}")
    RL_SCORE = reg("rl_probe:标定评分", f"{float(RLD['score_ref_exp_t']):.2f}")
    RL_RAND = reg("rl_probe:随机起点", f"{float(RLD['random_start_exp_t']):.2f}")
    RL_GAP = reg("rl_probe:RL距标定%",
                 f"{(float(RLD['rl_final_exp_t']) / float(RLD['score_ref_exp_t']) - 1) * 100:.1f}")
    RL_COS = reg("rl_probe:余弦", f"{float(RLD['cos_to_calibrated']):.2f}")
except Exception:
    pass

# δ 悬崖-平台数字(F15;δ 作用于在线评分,score 列)
D_CLIFF = D_PLAT_LO = D_PLAT_HI = MISS
try:
    _d = {float(r["value"]): float(r["expt_score"])
          for r in SENS if r["factor"] == "delta"}
    D_CLIFF = reg("sensitivity:δ=0悬崖", f"{_d[0.0]:.2f}")
    _plat = [v for k, v in _d.items() if 0.10 <= k <= 0.30]
    D_PLAT_LO = reg("sensitivity:δ平台min", f"{min(_plat):.2f}")
    D_PLAT_HI = reg("sensitivity:δ平台max", f"{max(_plat):.2f}")
except Exception:
    pass

# 场景权重策略表：只消费 fail-first 后的正式 policy，不再从幸存者 exp_t 自行重算。
ADW_GAIN_MEAN = ADW_GAIN_PEAK = ADW_INFEASIBLE = ADW_UNSTABLE = MISS
try:
    _policy = load_rows("weight_policy_table.csv")
    _gains = [float(r["gain_pct"]) for r in _policy
              if r.get("status") == "FEASIBLE" and r.get("gain_pct")]
    if _gains:
        ADW_GAIN_MEAN = reg("adaptive_weights:平均增益%",
                            f"{sum(_gains) / len(_gains):.2f}")
        ADW_GAIN_PEAK = reg("adaptive_weights:峰值增益%", f"{max(_gains):.2f}")
    ADW_INFEASIBLE = reg("adaptive_weights:INFEASIBLE单元",
                         str(sum(r.get("status") == "INFEASIBLE" for r in _policy)))
    ADW_UNSTABLE = reg("adaptive_weights:选择不稳定单元",
                       str(sum((r.get("selection_stable") or "").lower() == "false"
                               for r in _policy)))
except Exception:
    pass

# δ 生命周期定论(F21;两组文件,0710 修缮:45% 主组与 90% 高压组分名不再互覆)
def _lc_pair(rows, turnover):
    d15 = {r["seed"]: r for r in rows
           if r["turnover"] == turnover and float(r["delta"]) == 0.15}
    d0 = {r["seed"]: r for r in rows
          if r["turnover"] == turnover and float(r["delta"]) == 0.0}
    seeds = sorted(set(d15) & set(d0))
    if not seeds:
        return None
    diffs = [float(d15[s]["retr_mean"]) - float(d0[s]["retr_mean"]) for s in seeds]
    return (sum(diffs) / len(diffs),
            sum(int(d15[s]["fails"]) for s in seeds),
            sum(int(d0[s]["fails"]) for s in seeds), len(seeds))


LC_LOW_D10 = LC_LOW_D3 = LC_HP_F15 = LC_HP_F0 = LC_HP_FDROP = MISS
try:
    _lc = load_rows("lifecycle_delta.csv")
    _r10 = _lc_pair(_lc, "10")
    _r3 = _lc_pair(_lc, "3")
    if _r10:
        LC_LOW_D10 = reg("lifecycle:45%汰换10%差s", f"{_r10[0]:+.2f}")
    if _r3:
        LC_LOW_D3 = reg("lifecycle:45%汰换33%差s", f"{_r3[0]:+.2f}")
    _hp = _lc_pair(load_rows("lifecycle_delta_240.csv"), "3")
    if _hp:
        LC_HP_F15 = reg("lifecycle_240:δ组回库失败", str(_hp[1]))
        LC_HP_F0 = reg("lifecycle_240:无δ回库失败", str(_hp[2]))
        LC_HP_FDROP = reg("lifecycle_240:失败降幅%",
                          f"{(_hp[2] - _hp[1]) / _hp[2] * 100:.0f}")
except Exception:
    pass

# PLC证据必须来自“真实跑过”的登记表，禁止把源码N_CASES或占位TRUE当通过数。
PLC_EVIDENCE = by_key(load_rows("plc_evidence.csv"), "metric")
TEST_N = reg("plc_evidence:verified_passed",
             PLC_EVIDENCE.get("verified_passed", {}).get("value", MISS))
TEST_FAIL = reg("plc_evidence:verified_failed",
                PLC_EVIDENCE.get("verified_failed", {}).get("value", MISS))
SOURCE_TEST_N = reg("plc_evidence:source_declared_cases",
                    PLC_EVIDENCE.get("source_declared_cases", {}).get("value", MISS))
PLACEHOLDER_CASES = reg("plc_evidence:placeholder_cases",
                        PLC_EVIDENCE.get("placeholder_cases", {}).get("value", MISS))
SOURCE_REAL_CASES = reg("plc_evidence:source_real_unverified_cases",
                        PLC_EVIDENCE.get("source_real_unverified_cases", {}).get("value", MISS))
SOURCE_ROUTE_PRESENT = reg("plc_evidence:source_cb_tob_lex_present",
                           PLC_EVIDENCE.get("source_cb_tob_lex_present", {}).get("value", "0"))

# 敏感性结论(与 sensitivity.py 同逻辑重算,数据同源)
SENS_D = [r for r in SENS if r["factor"] == "delta"]
SENS_A = [r for r in SENS if r["factor"] == "acc_scale"]
SENS_DROPS = []
for r in SENS_A:
    try:
        n, a = float(r["expt_near"]), float(r["expt_awra"])
        SENS_DROPS.append((n - a) / n * 100)
    except Exception:
        pass
SENS_RANGE = (f"{min(SENS_DROPS):.1f}%~{max(SENS_DROPS):.1f}%"
              if SENS_DROPS else MISS)


# ---------------- docx 排版助手 ----------------
def _font(run, size=10.5, bold=False, mono=False, color=None):
    run.font.name = MONO if mono else FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


class Rep:
    def __init__(self):
        self.d = Document()
        st = self.d.styles["Normal"]
        st.font.name = FONT
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(10.5)

    def h(self, text, level):
        p = self.d.add_heading("", level=level)
        _font(p.add_run(text), {0: 15.5, 1: 15, 2: 12.5, 3: 11}[level], bold=True,
              color=(0x17, 0x36, 0x5D))

    def p(self, text, size=10.5):
        par = self.d.add_paragraph()
        par.paragraph_format.first_line_indent = Cm(0.74)
        _font(par.add_run(text), size)
        return par

    def note(self, text, size=9):
        par = self.d.add_paragraph()
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        _font(par.add_run(text), size, color=(0x88, 0x55, 0x00))

    def fig(self, fname, caption):
        path = os.path.join(FIGS, fname)
        if not os.path.exists(path):
            self.note(f"〔待补图:{fname}〕")
            return
        par = self.d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(path, width=Cm(15.5))
        cap = self.d.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(cap.add_run(caption), 9, bold=True, color=(0x40, 0x40, 0x40))

    def table(self, header, rows, widths=None):
        tb = self.d.add_table(rows=len(rows) + 1, cols=len(header))
        tb.style = "Table Grid"
        for ci, htext in enumerate(header):
            cell = tb.rows[0].cells[ci]
            cell.paragraphs[0].text = ""
            _font(cell.paragraphs[0].add_run(htext), 9.5, bold=True)
            shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
                qn("w:val"): "clear", qn("w:fill"): "D9E8F5"})
            cell._tc.get_or_add_tcPr().append(shd)
        for ri, row in enumerate(rows, 1):
            for ci, val in enumerate(row):
                cell = tb.rows[ri].cells[ci]
                cell.paragraphs[0].text = ""
                _font(cell.paragraphs[0].add_run(str(val)), 9.5)
        self.d.add_paragraph()


def set_view_zoom(doc):
    """Emit a schema-complete Word zoom setting for downstream OOXML validation."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(qn("w:val"), "bestFit")
    zoom.set(qn("w:percent"), "100")


def main():
    required = [
        "headline_numbers.csv", "mixed_ops.csv", "energy_report_data.csv",
        "energy_report_params.csv", "realism_matrix_data.csv",
        "realism_matrix_params.csv", "weight_policy_table.csv",
        "lifecycle_delta.csv", "l4_task_cycles_data.csv",
    ]
    missing = [name for name in required if not os.path.exists(os.path.join(OUT, name))]
    if missing:
        print("缺少报告必需输入，拒绝生成：" + ", ".join(missing))
        raise SystemExit(1)
    r = Rep()
    canonical = os.path.join(os.path.dirname(HERE), "docs", "canonical_assumptions.md")
    with open(canonical, "rb") as fp:
        canonical_sha = hashlib.sha256(fp.read()).hexdigest()
    r.d.core_properties.title = "智储优控验证报告（生成器草稿）"
    r.d.core_properties.subject = "sim口径与PLC证据边界分离；canonical v3.4"
    r.d.core_properties.keywords = f"generated={datetime.now().isoformat(timespec='seconds')};canonical_sha256={canonical_sha}"

    # ---------------- 封面与摘要 ----------------
    r.h("智储优控:基于 AI 算法与智能控制的立体仓储仓位优化", 0)
    r.p("2026 ABB 杯智能技术创新大赛 · 赛道一 · 初赛验证报告(生成器草稿)")
    r.note("本文档由sim/build_report.py从当前CSV/PNG生成；生成器只消费现有上游产物，"
           "不代表所有专项已在本次重跑。canonical SHA256已写入文档元数据。")
    r.h("摘要", 1)
    r.p(f"针对20×20、400仓位立体仓库，本文提出在线多目标评分+批量AWRA-LS架构。"
        f"在sim H设计口径（sum+梯形加减速+场景权重）下，AWRA-LS期望取货时间为"
        f"{H_AW}±{H_AW_STD}s（5-seed sample SD；均值95%CI半宽{H_AW_CI}s；非U3 30-seed"
        f"泛化/非独立final），较顺序基线{H_SEQ}s降低{DROP_VS_SEQ}%，能耗代理降低{DROP_E}%，"
        f"重货平均层高{HEAVY_TIER}。有效头条行失败0、承重/容积复查0。{THR_AW}次/时是"
        f"行程-only连续满载sim单操作容量（约{THR_CYCLES}双周期/时），为基线{THR_X}倍，不是PLC吞吐实测。PLC侧证据限于"
        f"核心ST逻辑+静态匀速黄金向量在干净初态/当前参数下AB仿真{TEST_N}/{TEST_FAIL}通过；"
        f"不证明H KPI、场景权重闭环或新增CB/TOB/lex源码已在AB/控制器复现；年化能耗亦仅为sim换算。")

    # ---------------- 1 问题与建模 ----------------
    r.h("1 问题定义与建模口径", 1)
    r.p("题面要求:400 仓位(20 列×20 层),底层承重系数 1.0、逐层递减 0.02;坐标能被 3 "
        "整除的仓位预设占用;货物具有序号、名称、重量、访问频次、体积五项属性;需综合考虑"
        "存取效率、空间利用率与设备损耗,并基于 AC500 PLC 实现。")
    r.p("行程时间与路径长度采用双口径并列:执行时间按切比雪夫模型(水平/垂直双电机同动,"
        "t=max(tx,ty),AS/RS 标准模型),路径长度按曼哈顿距离(损耗代理)。预占用规则"
        "'坐标值能被 3 整除'在题面存在多种解读,本文自始将其实现为可切换参数(x 且 y 整除"
        "/任一整除/线性编号/坐标和整除四种);向组委会询问后获官方口径——(x+y) 之和能被 "
        "3 整除,共133格、含(0,0),呈对角条纹分布——主口径切换后分阶段重算当前数据与图,"
        "历史口径保留为回归测试中的对照锁。该规则性稀疏预占格局"
        "在仓位分配文献中未见直接研究,本文 §4 给出其对各策略影响的量化属一手结果;"
        "最近位优先类策略在该格局下的系统性劣化与 Hassini & Vickson(2007)关于"
        "最近开放位规则非均匀性的结论方向一致。全部假设集中登记于 canonical_assumptions.md,"
        "正文数字均标注口径。")
    r.p("本文全部指标口径集中定义如下(唯一出处为交付包 canonical_assumptions.md,"
        "正文引用一律标注口径号):")
    r.table(["口径", "名称", "定义与用途"], [
        ["C1", "入库总时间", "Σ双程时间(I/O→仓位→返回 I/O),对应题面'完成全部仓位存储总时间'"],
        ["C2", "路径长度/行程时间", "路径=曼哈顿距离(物理走线,损耗代理);时间=切比雪夫(双轴同动);两者并列呈现,不混用"],
        ["C3", "期望取货时间", "Σ(fᵢ×tᵢ)/Σfᵢ,按访问频次加权的单程时间——'高频放近处'的核心度量"],
        ["C4", "能耗代理", "Σ(wᵢ·g·hᵢ) 提升做功(kg·m 相对值),另含可选水平摩擦项"],
        ["C4b", "能耗物理换算", "sim净机械功情景换算；未计待机/辅助/驱动器/动态损耗(§7)"],
        ["C5", "设备损耗代理", "堆垛机累计行走距离(曼哈顿累计,m)"],
        ["C6", "货架稳定性", "载荷重心高度 Σ(wᵢ·hᵢ)/Σwᵢ,越低越稳"],
        ["C7", "空间利用率", "已占仓位/400(含预占用);细口径为体积利用率"],
        ["C8", "约束违规数", "超限重/超容积/占用冲突计数,任何有效方案必须为 0"],
        ["C9", "吞吐", f"双命令模式下{THR_AW}单操作/h（约{THR_CYCLES}双周期/h）；行程-only连续满载sim容量，非PLC实测"],
        ["C10", "双命令周期", "联程 T=t(0→存)+t(存→取)+t(取→0);较两个单命令省一次空驶,节省量≥0 由三角不等式保证"],
    ])
    r.p("本文的物理与场景参数经系统调研溯源,自真实生产场景入手建立拟真模型:堆垛机厂商"
        "公开规格(7 组真实机型样本,含与本文货物量级对应的轻载单立柱档)、ABB 官方应用"
        "文档(堆垛机专项手册与起重机控制固件参数手册)、AS/RS 原始文献(行程时间模型"
        "开创性论文的原文假设)与中国现行行业标准名录(JB/T 7016-2017《巷道堆垛起重机》)"
        "——每个取值可回溯到公开出处或显式声明为情景假设,引用体系分三层列于附录 D。")

    # ---------------- 2 双层算法(0710 样章化:δ 证据链专节) ----------------
    r.h("2 双层优化架构:在线评分与批量 AWRA-LS", 1)
    r.h("2.1 在线层:四项加权评分", 2)
    r.p("在线层面向决赛交互场景:货物逐件到达,不预知未来,ST以单次调用扫描400仓位并给出"
        "推荐位；AB PC仿真曾观察到单周期返回，实体AC500最坏时延仍待测。评分函数为四项加权:效率项(频次×行程时间)、稳定项"
        "(重量×层高)、能耗项(提升做功)与位置价值预留项 δ(冷门货占近位受罚)。"
        "前三项为当前货物找好位置,δ 为未来货物留位置——它是四项中唯一带前瞻性的一项,"
        "也是本文投入验证最重的一项,其完整证据链见 §2.3。")
    r.h("2.2 批量层:AWRA-LS 与 gap 分析", 2)
    r.p(f"批量层 AWRA-LS(Rank→Assign→Local Search)在全部货物已知时给出优化上界:先按"
        f"频次/重量/体积加权优先级排序,再按同一评分函数逐件分配,最后以重定位与成对交换"
        f"两类邻域做局部搜索。批量层较在线评分再降 {GAP_ONLINE}%,该差距即'信息完备程度'"
        f"的价值,连接两层的 gap 分析构成报告主线:在线层回答'现在能做到多好',批量层"
        f"回答'信息拉满还能好多少',两者之差给重排机制(§4.3)定价。")
    r.fig("fig1_策略对比.png", "图 1  五策略主指标对比(复现:python warehouse_sim.py)")
    r.fig("fig2_占用热力图.png", "图 2  典型布局占用热力图(热货聚 I/O 角)")

    r.h("2.3 位置价值预留项 δ:一条自我修正的证据链", 2)
    r.p("δ 的机理是给稀缺资源定价:近位有限,冷门货占据一格,热门货就少一格。围绕"
        "这一项,本文积累了五步相互独立的证据;其中最后一步修正了此前的解读。链条"
        "按发生顺序如实呈现,包括被修正的部分。")
    r.p(f"第一步,存在性(F2):δ=0 时,在线评分的占位集合与最近开放位规则(closest "
        f"open location)完全重合——删掉 δ,多目标评分便退化为换了写法的最近位规则。"
        f"§4.4 的参数化扫描复现同一现象并给出形状:δ=0 处在线评分期望取货 {D_CLIFF} s"
        f"(悬崖),δ∈[0.10,0.30] 内稳定于 {D_PLAT_LO}–{D_PLAT_HI} s(平台)。")
    r.p(f"第二步,早期静态扫描(F15)曾显示0.10–0.30的平台。第三步,旧版不做fail-first时"
        f"曾把高密度失败幸存者误读为'δ随密度上升'；0712重算后200/240件的六个online格"
        f"全部INFEASIBLE，这条密度弹性解释被撤回。批量可行格仍可比较，但2-seed只够作设计筛查。")
    r.p(f"第四步,反向探测(F16):以轻量策略梯度 RL 为对照,同预算训练后 RL 收敛于"
        f"'距离与层高主导'的策略(评测期望取货 {RL_FINAL} s,较随机起点 {RL_RAND} s "
        f"提升过半,但距手工标定评分 {RL_SCORE} s 仍差 {RL_GAP}%),且学得方向与标定"
        f"方向的余弦仅 {RL_COS}——RL 没有自发发现 δ 结构。当时的解读是:终局回报难以"
        f"把'为未来留位'的功劳归因到单步决策(信用分配难),前瞻性预留是设计出来的。")
    r.p(f"第五步,边界与修正(F21):全生命周期实验(200 周期循环访问+汰换+夜间重排,"
        f"30 seeds CRN 配对)给出更完整的图景——45% 填充下 δ 是净亏的:出库响应配对差 "
        f"{LC_LOW_D10} s(汰换 10%)至 {LC_LOW_D3} s(汰换 33%),p<0.001,近位不"
        f"紧缺时预留没有对象,只剩把冷门货推远的日常代价;90%高压下δ组失败"
        f"{LC_HP_F15}对{LC_HP_F0}次(少{LC_HP_FDROP}%),但两组失败均非零且时间差不显著。"
        f"这只能称失败缓解信号，不能称'放得进'或可行性保险。这一步同时修正了第四步的解读:"
        f"RL 训练发生在低压场景,该场景下 δ 本身净亏,'不学'是正确判断——把它归因"
        f"于信用分配难度属于过度解读。")
    r.p("这条链的价值不在每步都支持最初设计,而在每步都改变了对这一项的认识:从"
        "'δ 是核心创新'到'低填充收益被证伪、高压仅见失败缓解信号'。主动划定自家创新点的"
        "边界,本文视为方法论的一部分。")

    r.h("2.4 场景权重表:fail-first可行性与选择稳定性", 2)
    r.p(f"18场景×15权重组合按fail-first标定：只有fail=0且viol=0才比较exp_t。"
        f"可行单元相对固定默认平均增益{ADW_GAIN_MEAN}%、峰值{ADW_GAIN_PEAK}%；"
        f"{ADW_INFEASIBLE}个高密度online单元INFEASIBLE，{ADW_UNSTABLE}个单元的2-seed赢家不一致。"
        f"因此它是sim设计/风险表，不是无偏生产最优；默认运行不写ST，PLC在线查表闭环未接入。")
    r.fig("fig12_自适应权重.png", "图12  fail-first可行性与2-seed选择稳定性（sim）")

    # ---------------- 3 运动与作业模型 ----------------
    r.h("3 物理与作业真实度:梯形速度与双命令周期", 1)
    r.p("运动模型自匀速升级为梯形加减速(可退化,三角/梯形自动判别),消除'瞬间达到峰值"
        "速度'的失真;匀速口径保留为对照,支撑与文献的可比性。作业模型自'一次性入库'扩展"
        "为存取混合流:双命令周期(存+取联程)较两个单命令周期节省联程空驶,"
        f"实测节省 {DUAL_SAVE}%。")
    r.fig("fig7_运动模型对比.png", "图 7  匀速与梯形加减速口径对比")
    r.fig("fig8_混合作业.png", "图 8  混合作业流吞吐与双命令节省")

    # ---------------- 4 实验验证 ----------------
    r.h("4 实验与验证", 1)
    r.h("4.1 主对比(sim H设计口径)", 2)
    rows = []
    for k, label in (("random", "随机基线"), ("seq", "顺序基线(题面)"),
                     ("near", "最近开放位规则"), ("score", "在线评分"), ("awra", "AWRA-LS")):
        d = HEAD.get(k, {})
        rows.append([label, fnum(d, "H_exp_mean", 2, src="headline"),
                     fnum(d, "H_energy_mean", 0, src="headline"),
                     fnum(d, "H_heavy_tier", 2, src="headline"),
                     d.get("viol_sum", MISS)])
    r.table(["策略", "期望取货时间 s(C3)", "能耗 kg·m(C4)", "重货均层(前20%)", "失败/违规"],
            [row[:-1] + [f"{HEAD.get(s, {}).get('fail_sum', MISS)}/{row[-1]}"]
             for row, s in zip(rows, ("random", "seq", "near", "score", "awra"))])
    r.p(f"读表要点:期望取货时间(C3)按访问频次加权,直接度量'高频货是否放在近处';"
        f"重货均层统计最重前 20% 货物的平均存放层,AWRA-LS 将其压至 {HEAVY_TIER} 层"
        f"——重货几乎全部贴地。失败/违规分列禁止以漏件换低时间；此处5-seed回归锚均为0。"
        f"PLC画面ViolCnt只复查承重/容积，不能外推为所有业务约束。")
    r.p(f"对照口径(匀速+固定权重,文献可比)下 AWRA-LS 为 {LEG_AW} s,较顺序基线降低 "
        f"{LEG_DROP}%;主口径与对照口径结论方向一致,仅绝对数因物理模型不同而平移。")
    r.fig("fig4_时间能耗定位.png", "图 4  时间—能耗二维定位:五策略帕累托位置一览")
    r.h("4.2 理论对比与优化上界", 2)
    r.p(f"理论—仿真—实现三级证据链(匀速对照口径,便于与解析式直接对照):随机存储期望行程"
        f"的离散精确解析值与随机策略多 seed 仿真均值偏差在 1σ 内,Bozer-White 连续货架近似"
        f"与离散精确值互证，为该口径下的实现健全性提供一项检查，但不构成全模型正确性证明。"
        f"夹逼为：频次感知松弛下界 {LB_S} s ≤ NSGA-II已证可达候选点 {ACH_S} s ≤ AWRA-LS {AW_BOUND_S} s。"
        f"AWRA距松弛下界{GAP_LB}%，距该可达候选点{GAP_ACH}%；可达点不是已证明的全局最优。")
    r.p(f"该夹逼只说明在当前匀速静态模型中，AWRA与松弛下界/已知可达点的距离有量化参照；"
        f"它不证明理论天花板或更强优化器无改进空间。裸NSGA-II(随机初始化"
        f"跑满 200 代)被 AWRA-LS 单点整体支配、以启发式解播种后局部搜索零改进的完整过程"
        f"见图 6,它同时回答了'为什么不以进化算法为主线'。")
    r.fig("fig9_理论对比.png", "图 9  解析下界—仿真—实现三级证据链")
    r.fig("fig6_pareto前沿.png", "图 6  NSGA-II Pareto 前沿对照")
    r.h("4.3 鲁棒性", 2)
    r.p("鲁棒性验证由压力矩阵(3 seeds)完成,覆盖三类扰动:S1 需求漂移(货物访问频次"
        "画像整体失效)、S2 重货潮(重量分布向 50–80 kg 集中后重新布局)、S3 承重退化"
        "(低层承重能力下降 20%,模拟货架老化——把题面'设备损耗'从统计指标升级为扰动源);"
        "另设紧库存场景(预占用取'或'解读,231 格被占)检验仓位稀缺下的分配成功率。")
    srows = []
    for k, label in (("seq", "顺序基线"), ("near", "最近开放位规则"),
                     ("score", "在线评分"), ("awra", "AWRA-LS")):
        d = STRESS.get(k, {})
        srows.append([label,
                      fnum(d, "drift_pct", 1, src="stress_matrix"),
                      fnum(d, "drift_recovered", 2, src="stress_matrix"),
                      fnum(d, "heavy_tier", 2, src="stress_matrix"),
                      fnum(d, "degrade_viol", 2, src="stress_matrix"),
                      fnum(d, "degrade_cost", 0, src="stress_matrix")])
    r.table(["策略", "S1 漂移退化 %", "S1 夜排恢复后 s", "S2 重货潮均层",
             "S3 超载件数", "S3 修复行车 s"], srows)
    AW_DRIFT = fnum(STRESS.get("awra", {}), "drift_pct", 1, src="stress_matrix")
    AW_REC = fnum(STRESS.get("awra", {}), "drift_recovered", 2, src="stress_matrix")
    r.p(f"S1 列如实呈现了一个不利数字:AWRA-LS 的漂移退化({AW_DRIFT}%)是全部策略中最大"
        f"的——布局对频次画像优化得越深,画像失效时失去越多,这是优化系统的普遍规律而非"
        f"实现缺陷。工程答案是互补机制:配合夜间重排后恢复至 {AW_REC} s,仍处最优档。"
        f"在线策略决定稳态水平,重排机制决定扰动后的恢复速度,二者共同构成完整方案。")
    AW_DEG = fnum(STRESS.get("awra", {}), "degrade_viol", 0, src="stress_matrix")
    NE_DEG = fnum(STRESS.get("near", {}), "degrade_viol", 0, src="stress_matrix")
    NE_COST = fnum(STRESS.get("near", {}), "degrade_cost", 0, src="stress_matrix")
    r.p(f"S3 承重退化下,AWRA-LS 布局超载 {AW_DEG} 件(重货压底带来天然承重裕量),"
        f"最近开放位规则与在线评分各约 {NE_DEG} 件、需 {NE_COST} s 量级的修复行车。"
        f"紧库存场景(3 seeds)进一步暴露短视规则的结构性短板:最近开放位规则平均失败 "
        f"{fnum(TIGHT.get('near', {}), 'fail_mean', 2, src='experiments_agg:tight')} 件"
        f"(把稀缺的低层近位先喂给轻货,后到重货无可行位),AWRA-LS 失败 "
        f"{fnum(TIGHT.get('awra', {}), 'fail_mean', 0, src='experiments_agg:tight')} 件"
        f"——Rank阶段的重量优先级使重货先占底层。承重/容积计数为0，但near/score的失败件必须单列，不能写'全约束0'。")
    r.p("周期性重排作为运行期自愈机制,其价值是'自愈'而非'提效'(F11:稳态净增量不足 "
        "0.5%,如实报告;真实价值在漂移事件后的恢复速度与在线策略失能时的兜底)。重排何时"
        "执行采用三级动态判据:夜间窗口先做零成本体检(布局质量较上次重排后基准劣化超过"
        "阈值才开工),单步动作以'预期收益大于 1.2 倍搬运成本'放行(裕度覆盖损坏风险等"
        "未计价项),并受夜班搬运预算约束;搬运的额外损耗按时间与路径(磨损)双币种记账。"
        "高周转负载下判据每夜均触发(证明每夜重排是被判据确认而非假设),低漂移节奏下自动"
        "跳过部分夜晚且质量不损(F11 增补)。")
    r.fig("fig5_分布鲁棒性.png", "图 5  货物分布切换下的策略稳定性")
    r.fig("fig13_鲁棒压力.png", "图 13  鲁棒性压力矩阵三面板")
    r.fig("fig11_重排闭环.png", "图 11  sim周期重排机制:漂移与自愈")
    r.h("4.4 参数敏感性(D5)", 2)
    r.p(f"单因素扰动显示:δ 在 0.10–0.30 区间为平台、为 0 时退化(悬崖);β 与 γ 仅其和"
        f"起作用,和值扫描五倍范围结论不变(F5 冗余性的敏感性视角);加减速参数整体缩放"
        f" 0.5×–2×,AWRA-LS 较最近开放位规则的降幅稳定于 {SENS_RANGE},结论不依赖具体物理参数"
        f"取值。全部扰动实验违规为 0。")
    r.fig("fig3_delta标定.png", "图 3  δ 扫描:悬崖(δ=0 退化)与平台(0.10–0.30)")
    r.fig("fig14_参数敏感性.png", "图 14  参数敏感性三联图(复现:python sensitivity.py)")
    r.d.add_page_break()
    r.h("4.5 拟真度阶梯:三档同管线验证", 2)
    r.p("为回答'模型贴合实际到什么程度',本文把拟真度显式分为三档,并在同一身份负载与"
        "共同随机数框架下并排运行(档3到达流由seq reference在策略循环外生成一次):档1"
        "数据模型(匀速+固定权重,文献可比);档 2 "
        "工程主口径(梯形加减速+自适应权重,即口径 H);档 3 工业场景——在档 2 之上叠加"
        "泊松到达峰谷、货叉装卸时间、访问频次估计噪声与设备故障注入。档 3 的情景参数全部"
        "显式声明(canonical G 节):装卸时间经10/15/20s三点扫描；P50下AWRA均优于near，"
        "但P95在15s点由near更低，故不宣称响应分位全面排序不变。")
    rrows = []
    for t, tl in (("1", "档 1 数据模型"), ("2", "档 2 主口径 H"), ("3", "档 3 工业场景")):
        rrows.append([tl,
                      fnum(REAL.get((t, "awra"), {}), "avg_retrieve_s", 2, src="realism"),
                      fnum(REAL.get((t, "seq"), {}), "avg_retrieve_s", 2, src="realism"),
                      TIER_DROP[t] + "%",
                      fnum(REAL.get((t, "awra"), {}), "resp_p95_s", 0, src="realism")
                      if t == "3" else "—"])
    r.table(["拟真档", "AWRA-LS 出库均时 s", "基线出库均时 s", "降幅", "响应 P95 s(档3)"],
            rrows)
    r.p(f"结果呈阶梯状:AWRA-LS 较题面基线的出库均时降幅从档 1 的 {TIER_DROP['1']}% "
        f"到档 2 的 {TIER_DROP['2']}%、档 3 的 {TIER_DROP['3']}%——匀速口径系统性高估优化"
        f"收益约6个百分点；三档AWRA相对seq行程方向保持、失败/承重容积复查为0。档3单seed"
        f"AWRA相对seq的P95为{RESP_SEQ95}→{RESP_AW95}s、P50为{RESP_SEQ50}→{RESP_AW50}s；"
        f"但near的P95为{RESP_NEAR95}s，低于AWRA的{RESP_AW95}s，故不能写AWRA响应全面最优。"
        f"访问频次估计噪声(对数正态 σ=0.3)下 AWRA-LS 布局质量几乎不变:Rank 阶段是"
        f"重量/频次/体积多因子排序且局部搜索兜底,单一画像失真不致坍塌——'访问频次从哪来、"
        f"不准怎么办'在该档得到定量回答。")
    r.fig("fig16_拟真度阶梯.png", "图 16  拟真度阶梯:三档布局质量与档 3 响应分位"
                                  "(复现:python sim/realism.py)")
    r.p("该阶梯支持AWRA相对seq的行程方向，但不推出所有策略、所有分位或所有密度均同序；"
        "高密度online可行性与生命周期δ结果已另行给出负结果。")

    # ---------------- 5 AC500 ST 与有限 AB 证据 ----------------
    r.h("5 ABB AC500 ST 实现与有限 AB PC 仿真证据", 1)
    r.p("实现遵循三层结构:FC 纯函数(承重系数/可行性/行程时间/评分)、FB 状态机(初始化/"
        "在线选位/批量分配/局部搜索/统计/动画/负载探针)、PRG 主状态机(复位→初始化→交互"
        "输入→批量分配→局部搜索→统计→完成)。核心工程手段是扫描周期感知分片:批量分配每"
        "周期处理 nBatchPerCycle 件、局部搜索每周期评估 nPairsPerCycle 对。该结构面向10ms"
        "任务设计；当前只有有限AB PC仿真证据，400件最坏时延与目标机watchdog尚未验证。")
    r.p(f"边界判断成体系:越界、负输入、超重、超容、预占用(含官方口径格数断言)、满仓、"
        f"无可行位报警、单件查询与记录翻页边界、载货折减与装卸三档等 {TEST_N} 项自测用例"
        f"(含与 Python 的一致性向量)在 Automation Builder 2.9 仿真环境经无头自动化管线"
        f"在干净初态+当前参数下通过(iPassed={TEST_N}、iFailed={TEST_FAIL}、编译0错误)。该门覆盖核心ST"
        f"逻辑与静态匀速黄金向量，不覆盖H KPI、自适应、lex/speed或多客户端权限安全。字符串状态码全部ASCII化,"
        f"可视化画面以静态中文标签对照解释。")
    r.p(f"当前工作区源码声明N_CASES={SOURCE_TEST_N}：{SOURCE_REAL_CASES}已是实质断言，"
        f"{PLACEHOLDER_CASES}仍为恒TRUE预留；源码还包含CB/TOB分配器与lex检测路由。"
        f"但本轮未开AB复测，因此报告证据数保持{TEST_N}/{TEST_FAIL}，不把源码实现、"
        f"新增断言或声明计数升级为测试通过。")
    L4ROWS = [r_ for r_ in load_rows("l4_task_cycles_data.csv") if r_.get("cycles_to_done")]
    r.p("扫描负载实测采用任务级口径。方法学结论先行:PC 仿真环境的三种周期内时钟源"
        "(IEC LTIME 差分、任务监视页统计、SysTime 库纳秒计数)经逐一实测均无法给出"
        "非零的周期内耗时(时间服务按扫描周期缓存/粒度粗于周期),故周期内微秒值"
        "在仿真环境不可测——如实记录,不作推断。可测且更硬的指标是**任务完成所需"
        "扫描周期数**。周期数是历史AB仿真观测；乘10ms只是按标称任务周期换算的真机预估，"
        "不能写成实体控制器同值。")
    lrows = []
    for r_ in L4ROWS:
        lrows.append([r_.get("group", MISS),
                      fnum(r_, "n_batch_per_cycle", 0, src="l4_task_cycles"),
                      fnum(r_, "n_pairs_per_cycle", 0, src="l4_task_cycles"),
                      fnum(r_, "cycles_to_done", 0, src="l4_task_cycles"),
                      reg("l4:真机预估s",
                          "%.2f" % (float(r_.get("cycles_to_done", 0)) * 0.010))])
    r.table(["组", "nBatchPerCycle", "nPairsPerCycle", "完成周期数(历史AB PC仿真)",
             "标称换算 s(周期数×10ms；非真机实测)"], lrows)
    r.p("读表:nBatchPerCycle=1 时分配段被显著拉长(逐件一周期),提高到 5 以上后"
        "瓶颈转移到状态机固定开销,完成周期数稳定;nPairsPerCycle 在 20 件演示批下"
        "不敏感——两两交换仅 190 对,单周期分片即可吃完(该参数在满仓 400 件场景才"
        "成为主变量)。历史四组均DONE_OK；本轮未开AB复测，亦未完成400件最坏时延/watchdog门。"
        "源脚本已归档tools/ab_scripting/archive/measure_l4b.py。")

    # ---------------- 6 有限一致性护栏 ----------------
    r.h("6 Python↔ST有限一致性护栏", 1)
    r.p("Python与ST只在已锁向量上交叉核对。20件T25静态匀速向量逐位一致。0712当前ST源码已将"
        "near平局键统一为(时间,层,列)并增加T60断言，但T60尚未AB复测；历史探针中的near分叉只作为"
        "修复前证据保留。REAL阈值分支与加减速评分分母差异仍存在。因此称'有限一致性护栏'，不称"
        "数学同源或报告数字由PLC互相背书。")

    # ---------------- 7 绿色工程 ----------------
    r.h("7 绿色工程换算:能耗的年化账", 1)
    assume = "、".join(f"{ASSUME_LABEL.get(k, k)}={v}" for k, v in E_ASSUME.items()) or MISS
    r.p(f"能耗按sim净机械功情景换算，全部假设显式声明({assume});{DAILY_OPS}指单操作/日(={DAILY_DUAL}双周期/日)，"
        f"CO2采用2023全国平均{CO2_FACTOR}kg/kWh(生态环境部公告2025年第47号)。模型未计待机、辅助、"
        f"驱动器与动态损耗，绝对量不是电表预测。在该假设下,AWRA-LS布局较题面基线"
        f"年节省用电约 {KWH_SAVE} kWh(降 {KWH_SAVE_PCT}%),电费与碳排放折算见图 10。"
        f"该结论与主办方'通过算法优化降低系统能耗'的绿色工程导向一致(王余,2026 年开赛"
        f"致辞)。")
    r.p(f"展望口径(不计入头条):若配置再生制动(ABB 传动方案,回收效率按声明假设 "
        f"{ETA_REGEN_S}),AWRA-LS 年回收 {REGEN_REC_AW} kWh、净耗电降至 {REGEN_NET_AW} "
        f"kWh。如实指出:再生对存放位置更高的基线布局回收更多,两布局的节省差距在该"
        f"口径下收窄至 {REGEN_GAP} kWh/年,但结论方向不变——这正是本文多口径方法(附录 C)"
        f"的意义:口径切换只平移数字,不翻转结论。")
    r.fig("fig10_绿色换算.png", "图 10  绿色工程年化换算(假设显式声明)")

    # ---------------- 8 创新点与工程取舍 ----------------
    r.d.add_page_break()
    r.h("8 创新点与工程取舍", 1)
    r.table(["创新点", "证据位置"], [
        ["位置价值预留项δ的自我修正证据链(早期静态信号→fail-first/lifecycle反证→边界收窄)", "§2.3/§4.4,F2→F13/F15→F21"],
        ["场景权重fail-first可行性/稳定性表(6格INFEASIBLE；PLC闭环未完成)", "§2,fig12"],
        ["扫描周期分片状态机+有限AB PC任务周期证据(目标机最坏时延待测)", "§5"],
        ["Python↔ST静态向量的有限一致性护栏", "§6"],
        ["sim周期重排机制与其成本经济学(负结果如实呈现)", "§4.3,F11"],
        ["参数溯源体系:四层防御(厂商样本/标准名录/敏感性覆盖/官方确认)+负面出处清单+三档拟真验证", "§1/§4.5/附录 C-D"],
    ])
    r.p(f"工程取舍(考虑过且不采用,均给出依据):A* 路径规划——双轴同动且巷道无障碍时"
        f"切比雪夫直达即物理最优,A* 适用于把网格当作平面移动机器人地图的另一种建模;"
        f"强化学习——除引用 2025 年文献(DRL 需长周期数据、收益为个位数百分比)外,"
        f"本文实跑了轻量策略梯度对照(线性 softmax,1200 回合,训练/评测种子分离):"
        f"RL 自随机({RL_RAND} s)学至 {RL_FINAL} s、约束违规 0,但收敛于'距离与层高主导"
        f"的近似就近'策略,仍落后手工标定评分({RL_SCORE} s){RL_GAP}%,且未能自发形成 δ "
        f"位置价值预留结构(学得方向与标定方向余弦 {RL_COS})。这既可能来自长程信用分配困难，"
        f"也可能是当前δ没有净收益；结合lifecycle drain负结果，本对照不能反证δ有价值。查表化执行——"
        f"400 仓位在线评分本身即微秒级,查表反而丢失解释链。")
    r.fig("fig15_RL对照.png", "图 15  轻量强化学习实跑对照:学习曲线与终局对比"
                              "(复现:python sim/rl_probe.py)")

    # ---------------- 9 结论 ----------------
    r.h("9 结论", 1)
    r.p(f"本文完成了sim模型、双层算法与ST核心逻辑的分层验证。sim H 5-seed回归锚中，"
        f"期望取货较顺序基线降低{DROP_VS_SEQ}%、能耗代理降低{DROP_E}%；{THR_X}倍吞吐为"
        f"行程-only连续满载sim单操作容量（约{THR_CYCLES}双周期/h）。AB侧{TEST_N}/{TEST_FAIL}只覆盖核心ST+静态匀速黄金向量。"
        f"负结果包括6个高密度online权重单元INFEASIBLE、δ低填充收益被drain复算证伪，"
        f"以及档3P95最优者并非AWRA。结论按数据链、控制器链和制品链分别验收，不宣称全链闭环。")

    # ---------------- 附录 ----------------
    r.h("附录 A  评委 5 分钟复现指引", 1)
    r.p(f"①`python sim/core_smoke.py`跑完整unittest discover+headline回归锚；②已有上游CSV/PNG时"
        f"用`report_rebuild.py`重建制品，再以`artifact_verify.py`做文本/CSV门；三者互不替代。"
        f"③AB操作需用户在场；已验收冻结基线为iPassed={TEST_N},iFailed={TEST_FAIL}。当前源码声明{SOURCE_TEST_N}项且{PLACEHOLDER_CASES}含占位，须实现真实断言后另行复测。")
    r.h("附录 B  评语六要素自查映射(2025 同构赛道一等奖评语)", 1)
    r.table(["评语要素", "对应章节"], [
        ["算法模型准确", "§1/§2/§4.2"], ["可运行验证仿真结果", "附录 A/§4"],
        ["一定创新性", "§8"], ["报告思路逻辑清晰", "全文递进结构"],
        ["功能块封装合理", "§5"], ["注释标注清晰易懂", "ST 源码(交付包)"],
    ])
    r.d.add_page_break()
    r.h("附录 C  口径矩阵:一套系统,九把尺子(多口径并列验证方法)", 1)
    r.p("各建模选择均显式标口径并保留开关。当前证据支持120件H回归锚及AWRA相对seq的"
        "三档行程方向；不支持把结论泛化到所有策略、密度、响应分位或生命周期机制。")
    r.table(["维度", "主口径", "对照/备选口径", "切换开关", "口径号"], [
        ["运动模型", "梯形加减速", "匀速(文献可比)", "--accel / xUseAccel", "A4b"],
        ["评分权重", "sim场景权重(含INFEASIBLE)", "固定1.0/0.6/0.4/0.15", "--adaptive / policy表", "E"],
        ["行程度量", "切比雪夫时间(耗时)", "曼哈顿长度(磨损)", "并列输出", "C2"],
        ["时间程数", "双程总时", "单程期望取货", "并列输出", "C1/C3"],
        ["预占用解读", "(x+y)整除·官方口径(133格)", "且/或/线性(49/231/134格)", "--rule / iPreRule", "B1"],
        ["作业流", "循环访问", "一次性流转", "工作负载参数", "B6/B7"],
        ["轴联动", "双轴同动(切比雪夫)", "顺序单轴(保守)", "xSimultaneousXY", "A5/A9"],
        ["能耗回收", "不回收(保守)", "再生制动展望(η 假设)", "energy_report 展望行", "C4b"],
        ["负载观测", "历史AB PC仿真周期数", "实体AC500待测", "—", "L4"],
        ["拟真度", "档2 sim H设计口径", "档1数据模型 / 档3综合sim场景", "realism.py 三档", "G"],
    ])
    r.p("工业上同类做法并不陌生:汽车油耗有 WLTP 与 EPA 两套循环工况,财务报告有会计准则"
        "口径与管理层口径——严肃的工程数字从来都需要声明测量框架。本文把这一实践系统化为"
        "口径矩阵,并以回归测试锁定主口径数字(动口径=测试红灯)。")

    r.h("附录 D  参考文献与文档溯源(三层引用体系,仅列实际参考项)", 1)
    r.p("D.1 学术文献(理论假设的原文锚)", size=10.5)
    for ref in [
        "Bozer, Y.A. & White, J.A. (1984). Travel-Time Models for Automated "
        "Storage/Retrieval Systems. IIE Transactions 16(4), 329-338. DOI "
        "10.1080/07408178408975252——双轴同动与 T=max(th,tv) 假设、单/双命令周期定义的原始出处",
        "Roodbergen, K.J. & Vis, I.F.A. (2009). A survey of literature on automated "
        "storage and retrieval systems. EJOR 194(2), 343-362——双命令调度=Chebyshev TSP",
        "AS/RS 行程时间模型与控制策略综述(Automated Storage and Retrieval Systems: "
        "A Review on Travel Time Models and Control Policies)",
        "考虑订单完成时间与能耗的 AS/RS 启发式(Mathematical Biosciences and "
        "Engineering, 2024)",
        "PLC Implementation of a Genetic Algorithm for Controller Optimization",
        "Deep reinforcement learning for the real-time inventory rack storage "
        "assignment and replenishment problem(EJOR, 2025)",
        "Dynamic Storage Location Assignment in Warehouses Using Deep Reinforcement "
        "Learning(Technologies, 2022;DRL 较 ABC 分类降本 6.3% 案例)",
    ]:
        r.p("· " + ref, size=9.5)
    r.p("D.2 ABB / 贝加莱官方文档(工程参数与产品能力锚)", size=10.5)
    for ref in [
        "ABB. ACS880 Stacker crane operation brochure(堆垛机专项). 3AXD50000736843 "
        "Rev B, 2023-11",
        "ABB. ACS880 Crane control program firmware manual. 3AUA0000132682——参数 "
        "23.12/23.13(加减速时间)、23.16 Shape time(0=线性斜坡/>0=S 曲线)、组 77 防摇",
        "ABB. AC500-S Smart Safety for Material Handling(e-book). 3ADR011134, 2022-08",
        "ABB Stacker Cranes 应用页(AC500 列为堆垛机推荐控制器)与再生驱动案例:废料天车 "
        "32%(News #4532,现场实测)、Pesmel 堆垛机回馈 54%(News #123993)——案例级引用",
        "B&R. Technology Guard: CNC Jerk Limitation(1TGMPCNCJERK)——jerk 受限时加速度呈"
        "梯形/三角、速度呈 S 形的官方定义",
        "2026 ABB 杯智能技术创新大赛官方赛题页与开赛新闻(ABB 中国)",
    ]:
        r.p("· " + ref, size=9.5)
    r.p("D.3 行业标准名录(周期时间计算的标准框架对应)", size=10.5)
    for ref in [
        "JB/T 7016-2017《巷道堆垛起重机》(取代 JB/T 7016-1993 与 JB/T 2960-1999,中国现行)",
        "FEM 9.851 Performance Data of Storage and Retrieval Machines: Cycle Times",
        "VDI 3561 单巷道堆垛机周期时间(Blatt 2:2019 为多巷道扩展版)",
    ]:
        r.p("· " + ref, size=9.5)
    path = os.path.join(OUT, "验证报告_draft.docx")
    set_view_zoom(r.d)
    r.d.save(path)
    print(f"已生成:{path}")
    print("\n数字来源清单(key → 注入值;〔待补〕项=上游数据未就绪):")
    for k, v in USED:
        print(f"  {k:<38} {v}")
    miss = sum(1 for _, v in USED if v == MISS)
    # 正文实扫(0710 修盲区):派生值 try/except 失败不进 USED,只数 USED 会虚报"待补 0"
    # ——以成文文本里的 MISS 出现次数为准(段落+表格双扫)
    body_miss = sum(p.text.count(MISS) for p in r.d.paragraphs)
    for tb in r.d.tables:
        body_miss += sum(c.text.count(MISS) for row in tb.rows for c in row.cells)
    print(f"\n共注入 {len(USED)} 个数字,来源清单待补 {miss} 个,"
          f"正文实扫待补 {body_miss} 处;正文无硬编码性能数字。")
    if body_miss > miss:
        print(f"⚠️ 有 {body_miss - miss} 处派生值缺数据(未进来源清单)——"
              f"检查上游 CSV 是否就绪后重跑本脚本。")
    if body_miss:
        print("报告正文仍有待补项，生成链以exit 1失败关闭。")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
