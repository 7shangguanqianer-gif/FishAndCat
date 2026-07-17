from docx import Document
p=r'F:\abb_wh_work\sim\out\验证报告_draft.docx'
d=Document(p)
text='\n'.join(x.text for x in d.paragraphs)+'\n'+'\n'.join(c.text for t in d.tables for row in t.rows for c in row.cells)
checks=['8.17±0.67','61.6%','30/30','Hausman','6 可视化仿真与人机界面','98.5%','75/0','21个交互件']
for x in checks: assert x in text,x
for bad in ['{LEX_ATTAIN}','〔待补〕','非U3 30-seed泛化','尚未AB复测','仍为恒TRUE预留']:
    assert bad not in text,bad
print('REPORT_TEXT_OK',len(d.paragraphs),len(d.tables),len(d.inline_shapes))
