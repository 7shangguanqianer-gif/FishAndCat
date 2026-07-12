# -*- coding: utf-8 -*-
"""FactoryIO仿真上手教程.md → Word 版(轻量 md 解析,0706)。
用法:python tools/make_factoryio_doc.py
输出:2_你要操作/FactoryIO仿真上手教程.docx
"""
import os
import re
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "2_你要操作", "FactoryIO仿真上手教程.md")
OUT = os.path.join(ROOT, "2_你要操作", "FactoryIO仿真上手教程.docx")

NAVY = RGBColor(0x17, 0x36, 0x5D)


def add_runs(p, text):
    """解析 **粗体** 与行内代码/URL,其余原样。"""
    for seg in re.split(r"(\*\*.+?\*\*|`[^`]+`|https?://\S+)", text):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = p.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith("`") and seg.endswith("`"):
            r = p.add_run(seg[1:-1]); r.font.name = "Consolas"
        elif seg.startswith("http"):
            r = p.add_run(seg); r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.underline = True; r.font.size = Pt(9)
        else:
            p.add_run(seg)


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "微软雅黑"; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")

    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    for ln in lines:
        s = ln.rstrip()
        if not s.strip():
            continue
        if s.startswith("> "):
            p = doc.add_paragraph(); add_runs(p, s[2:])
            p.runs and setattr(p.runs[0].font, "italic", True)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.size = Pt(9.5)
        elif s.startswith("# "):
            h = doc.add_heading(s[2:], level=0)
            for r in h.runs:
                r.font.color.rgb = NAVY
        elif s.startswith("## "):
            h = doc.add_heading(s[3:], level=1)
            for r in h.runs:
                r.font.color.rgb = NAVY
        elif re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", s))
        elif s.lstrip().startswith("- "):
            indent = len(s) - len(s.lstrip())
            p = doc.add_paragraph(style="List Bullet 2" if indent else "List Bullet")
            add_runs(p, s.lstrip()[2:])
        else:
            p = doc.add_paragraph(); add_runs(p, s)

    doc.save(OUT)
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
