# -*- coding: utf-8 -*-
"""0712 轨A：仅重生本批次指定的三份用户 DOCX，并做最小内容校验。"""
import os
import sys
import zipfile

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
sys.path.insert(0, ROOT)
from tools.make_user_docs import convert  # noqa: E402

USER = os.path.join(ROOT, "1_给你看")
NAMES = ("周报_第1期_0710", "项目通俗导览", "实验发现")

for name in NAMES:
    src = os.path.join(USER, name + ".md")
    dst = os.path.join(USER, name + ".docx")
    convert(src, dst)
    assert zipfile.is_zipfile(dst), dst
    doc = Document(dst)
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    for required in ("8.17±0.67", "30/30"):
        assert required in text, (dst, required)
    print(f"GENERATED_OK\t{name}.docx\t{os.path.getsize(dst)}")
