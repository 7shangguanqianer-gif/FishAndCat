# -*- coding: utf-8 -*-
"""给无Markdown源的孤儿DOCX加显眼历史标记；可重复运行且不重复插入。"""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PATH = os.path.join(ROOT, "1_给你看", "项目全景说明书.docx")
MARK = "【历史快照·已过时·勿引用当前数字】"


def style(run, size=14):
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def main():
    doc = Document(PATH)
    existing = "\n".join(p.text for p in doc.paragraphs[:5])
    if MARK not in existing:
        first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        p = first.insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style(p.add_run(MARK), 16)
        q = first.insert_paragraph_before()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = q.add_run("本文件无可再生Markdown源，含旧口径；当前事实以docs/canonical_assumptions.md及当前再生报告为准。")
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

    for section in doc.sections:
        header = section.header
        if MARK not in "\n".join(p.text for p in header.paragraphs):
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style(p.add_run(MARK), 9)

    props = doc.core_properties
    props.subject = "Historical snapshot; superseded"
    props.keywords = "historical_snapshot;superseded;do_not_quote"
    props.comments = (
        "historical_snapshot=true; no Markdown source; current truth is "
        "docs/canonical_assumptions.md and regenerated deliverables"
    )
    doc.save(PATH)
    print(f"已标记历史孤儿DOCX: {PATH}")


if __name__ == "__main__":
    main()
