# -*- coding: utf-8 -*-
r"""
make_user_docs.py — 用户层文档 md→docx 一键再生(用户偏好:md 难读,要 Word 版)
纪律:**md 是唯一的源,docx 是生成物**——改内容只改 md,然后重跑本脚本;
     doc_check.py 会核对"docx 不旧于 md",忘了重生成会亮红灯(收尾五联动第④步)。
范围:1_给你看\*.md 与 2_你要操作\*.md → 同目录同名 .docx。
支持的 md 子集(本项目文档全部在此子集内):
  # ## ### 标题 | 无序/有序列表 | 表格 | ``` 代码块 | > 引用 | **粗体** `行内代码` | --- 分隔线
运行:python tools/make_user_docs.py
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
USER_DIRS = ["1_给你看", "2_你要操作"]
FONT = "Microsoft YaHei"
MONO = "Consolas"


def set_font(run, size=10.5, bold=False, mono=False, color=None):
    run.font.name = MONO if mono else FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_rich(par, text, size=10.5, base_bold=False):
    """解析 **粗体** 与 `行内代码`。"""
    for piece in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            set_font(par.add_run(piece[2:-2]), size, bold=True)
        elif piece.startswith("`") and piece.endswith("`"):
            set_font(par.add_run(piece[1:-1]), size, mono=True, color=(0x1A, 0x4D, 0x7C))
        else:
            set_font(par.add_run(piece), size, bold=base_bold)


def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(10.5)


def convert(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    style_doc(doc)
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s.startswith("```"):                      # 代码块(灰底等宽,整块)
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            for code_ln in buf:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Pt(12)
                set_font(p.add_run(code_ln if code_ln else " "), 9, mono=True,
                         color=(0x33, 0x33, 0x33))
                shd = p._p.get_or_add_pPr().makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:fill"): "F2F2F2"})
                p._p.get_or_add_pPr().append(shd)
            doc.add_paragraph()
            continue
        if s.startswith("|") and i + 1 < len(lines) and \
                re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):   # 表格
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            rows.pop(1)                               # 去分隔行
            ncol = max(len(r) for r in rows)
            tb = doc.add_table(rows=len(rows), cols=ncol)
            tb.style = "Table Grid"
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, r in enumerate(rows):
                for ci in range(ncol):
                    cell = tb.rows[ri].cells[ci]
                    cell.paragraphs[0].text = ""
                    add_rich(cell.paragraphs[0], r[ci] if ci < len(r) else "",
                             size=9.5, base_bold=(ri == 0))
                    if ri == 0:
                        shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
                            qn("w:val"): "clear", qn("w:fill"): "D9E8F5"})
                        cell._tc.get_or_add_tcPr().append(shd)
            doc.add_paragraph()
            continue
        if s.startswith("#"):                         # 标题
            level = min(len(s) - len(s.lstrip("#")), 3)
            text = s.lstrip("#").strip()
            p = doc.add_heading("", level=level)
            run = p.add_run(text)
            set_font(run, {1: 16, 2: 13, 3: 11.5}[level], bold=True,
                     color=(0x17, 0x36, 0x5D))
            i += 1
            continue
        if s.startswith(">"):                         # 引用块
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            add_rich(p, s.lstrip("> ").strip(), size=10)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            i += 1
            continue
        if s in ("---", "***"):                       # 分隔线 → 空行
            doc.add_paragraph()
            i += 1
            continue
        m = re.match(r"^(\s*)[-*]\s+(.*)$", ln)       # 无序列表
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_rich(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)\d+[.)]\s+(.*)$", ln)    # 有序列表
        if m:
            p = doc.add_paragraph(style="List Number")
            add_rich(p, m.group(2))
            i += 1
            continue
        p = doc.add_paragraph()                       # 普通段落
        add_rich(p, s)
        i += 1
    doc.save(docx_path)


def main():
    made = []
    for d in USER_DIRS:
        full = os.path.join(ROOT, d)
        for f in sorted(os.listdir(full)):
            if f.endswith(".md"):
                src = os.path.join(full, f)
                dst = os.path.join(full, f[:-3] + ".docx")
                convert(src, dst)
                made.append(os.path.join(d, f[:-3] + ".docx"))
    print(f"已生成 {len(made)} 份 Word:")
    for m in made:
        print("  ", m)
    print("纪律:改内容只改 .md,改完重跑本脚本;doc_check 会核对 docx 不旧于 md。")


if __name__ == "__main__":
    main()
